"""
Analyze propeller logs, produce figures, and generate a DOCX research paper
author: siddarth boggarapu
"""

from pathlib import Path
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from datetime import datetime
import zipfile
import os

# Output folder
output_dir = Path("propeller_analysis_output")
output_dir.mkdir(exist_ok=True)

# Input files - update names if different
files = [
    Path("simulation_20251209_145334.csv")
]

# Read CSVs
dfs = []
for f in files:
    if not f.exists():
        print(f"WARNING: input file not found: {f}")
        continue
    try:
        df = pd.read_csv(f, parse_dates=["timestamp"])
    except Exception:
        df = pd.read_csv(f)
        if "timestamp" in df.columns:
            try:
                df["timestamp"] = pd.to_datetime(df["timestamp"])
            except Exception:
                pass
    df["source_file"] = f.name
    dfs.append(df)

if not dfs:
    raise SystemExit("No CSV inputs found. Place your CSVs next to this script or update 'files' list.")

df_all = pd.concat(dfs, ignore_index=True, sort=False)

# Coerce numeric columns
numeric_cols = ["rpm", "pitch_deg", "blades", "thrust_N", "power_W", "efficiency",
                "disk_loading_N_m2", "power_loading_N_W", "tip_speed_m_s", "tip_mach", "cl", "cd"]
for c in numeric_cols:
    if c in df_all.columns:
        df_all[c] = pd.to_numeric(df_all[c], errors="coerce")

# Sort by timestamp if present and add time index
if "timestamp" in df_all.columns and pd.api.types.is_datetime64_any_dtype(df_all["timestamp"]):
    df_all = df_all.sort_values("timestamp").reset_index(drop=True)
    t0 = df_all["timestamp"].iloc[0]
    df_all["t_s"] = (df_all["timestamp"] - t0).dt.total_seconds()
else:
    df_all["t_s"] = np.arange(len(df_all))

# Summary statistics
summary = df_all[numeric_cols].describe().transpose()
summary_csv = output_dir / "summary_stats.csv"
summary.to_csv(summary_csv)

# Helper to save figures
plots = []
def save_fig(fig, name):
    path = output_dir / name
    fig.savefig(path, bbox_inches="tight", dpi=150)
    plt.close(fig)
    plots.append(path)
    return path

# Time-series plots
if "thrust_N" in df_all.columns:
    fig = plt.figure(figsize=(8,4))
    plt.plot(df_all["t_s"], df_all["thrust_N"])
    plt.xlabel("Time (s)")
    plt.ylabel("Thrust (N)")
    plt.title("Thrust vs Time")
    save_fig(fig, "fig_thrust_time.png")

if "power_W" in df_all.columns:
    fig = plt.figure(figsize=(8,4))
    plt.plot(df_all["t_s"], df_all["power_W"])
    plt.xlabel("Time (s)")
    plt.ylabel("Power (W)")
    plt.title("Power vs Time")
    save_fig(fig, "fig_power_time.png")

if "efficiency" in df_all.columns:
    fig = plt.figure(figsize=(8,4))
    plt.plot(df_all["t_s"], df_all["efficiency"])
    plt.xlabel("Time (s)")
    plt.ylabel("Efficiency")
    plt.title("Efficiency vs Time")
    save_fig(fig, "fig_eff_time.png")

# Scatter plots vs RPM
if {"thrust_N","rpm"}.issubset(df_all.columns):
    fig = plt.figure(figsize=(6,4))
    plt.scatter(df_all["rpm"], df_all["thrust_N"])
    plt.xlabel("RPM")
    plt.ylabel("Thrust (N)")
    plt.title("Thrust vs RPM (scatter)")
    save_fig(fig, "fig_thrust_vs_rpm.png")

if {"power_W","rpm"}.issubset(df_all.columns):
    fig = plt.figure(figsize=(6,4))
    plt.scatter(df_all["rpm"], df_all["power_W"])
    plt.xlabel("RPM")
    plt.ylabel("Power (W)")
    plt.title("Power vs RPM (scatter)")
    save_fig(fig, "fig_power_vs_rpm.png")

# Histograms
if "thrust_N" in df_all.columns:
    fig = plt.figure(figsize=(6,4))
    plt.hist(df_all["thrust_N"].dropna(), bins=20)
    plt.xlabel("Thrust (N)")
    plt.title("Thrust distribution")
    save_fig(fig, "fig_thrust_hist.png")

if "power_W" in df_all.columns:
    fig = plt.figure(figsize=(6,4))
    plt.hist(df_all["power_W"].dropna(), bins=20)
    plt.xlabel("Power (W)")
    plt.title("Power distribution")
    save_fig(fig, "fig_power_hist.png")

if "tip_mach" in df_all.columns:
    fig = plt.figure(figsize=(6,4))
    plt.hist(df_all["tip_mach"].dropna(), bins=20)
    plt.xlabel("Tip Mach")
    plt.title("Tip Mach distribution")
    save_fig(fig, "fig_tipmach_hist.png")

# Boxplot by blades if blades vary
if "blades" in df_all.columns and df_all["blades"].nunique()>1 and "thrust_N" in df_all.columns:
    fig = plt.figure(figsize=(6,4))
    groups = [grp["thrust_N"].dropna().values for name,grp in df_all.groupby("blades")]
    labels = [str(name) for name,grp in df_all.groupby("blades")]
    plt.boxplot(groups, labels=labels)
    plt.xlabel("Blades")
    plt.ylabel("Thrust (N)")
    plt.title("Thrust by Blade Count (boxplot)")
    save_fig(fig, "fig_thrust_by_blades_box.png")

# Correlation matrix (simple imshow)
corr_cols = ["rpm","pitch_deg","thrust_N","power_W","efficiency","tip_mach"]
present = [c for c in corr_cols if c in df_all.columns]
if len(present)>=2:
    corr = df_all[present].corr()
    fig = plt.figure(figsize=(6,5))
    plt.imshow(corr, interpolation='nearest')
    plt.colorbar()
    plt.xticks(range(len(present)), present, rotation=45, ha='right')
    plt.yticks(range(len(present)), present)
    plt.title("Correlation matrix")
    save_fig(fig, "fig_correlation_matrix.png")

# Save cleaned combined CSV
clean_csv = output_dir / "combined_cleaned.csv"
df_all.to_csv(clean_csv, index=False)

# Create a DOCX research paper (requires python-docx)
docx_path = output_dir / "research_paper_siddarth_boggarapu.docx"
try:
    from docx import Document
    from docx.shared import Inches, Pt
    doc = Document()
    doc.add_paragraph("Boggarapu 1")
    doc.add_paragraph("siddarth boggarapu\nInstructor Name\nCourse Number\n" + datetime.now().strftime("%B %d, %Y"))

    title = doc.add_paragraph()
    title.alignment = 1
    run = title.add_run("Quantitative Analysis of Propeller Performance: Experimental Logs and BEMT Simulation")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph("")

    doc.add_heading("Abstract", level=2)
    doc.add_paragraph("This study analyses propeller performance logs... (short abstract)")

    doc.add_heading("Introduction", level=2)
    doc.add_paragraph("Propeller performance is critical...")

    doc.add_heading("Methods", level=2)
    doc.add_paragraph("Data from simulator CSVs...")

    doc.add_heading("Results", level=2)
    doc.add_paragraph("Key plots follow:")

    for fig_path in plots:
        doc.add_paragraph()
        doc.add_picture(str(fig_path), width=Inches(6))
        doc.add_paragraph(f"Fig. — {fig_path.name}", style='Caption')

    doc.add_heading("Summary statistics", level=3)
    stats_to_show = summary.reset_index().rename(columns={"index":"metric"})
    rows = min(len(stats_to_show), 10)
    table = doc.add_table(rows=1+rows, cols=6)
    hdr = table.rows[0].cells
    hdr[0].text = "metric"
    hdr[1].text = "count"
    hdr[2].text = "mean"
    hdr[3].text = "std"
    hdr[4].text = "min"
    hdr[5].text = "max"
    for i in range(rows):
        r = stats_to_show.iloc[i]
        cells = table.rows[i+1].cells
        cells[0].text = str(r["metric"])
        cells[1].text = f"{r.get('count', ''):.0f}" if not pd.isna(r.get('count', np.nan)) else ""
        cells[2].text = f"{r.get('mean', ''):.3f}" if not pd.isna(r.get('mean', np.nan)) else ""
        cells[3].text = f"{r.get('std', ''):.3f}" if not pd.isna(r.get('std', np.nan)) else ""
        cells[4].text = f"{r.get('min', ''):.3f}" if not pd.isna(r.get('min', np.nan)) else ""
        cells[5].text = f"{r.get('max', ''):.3f}" if not pd.isna(r.get('max', np.nan)) else ""

    doc.add_heading("Discussion", level=2)
    doc.add_paragraph("Discussion text...")

    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph("Conclusion text...")

    doc.add_heading("Works Cited", level=2)
    doc.add_paragraph("References available on request.")

    doc.save(docx_path)
    docx_created = True
except Exception as e:
    docx_created = False
    txt_path = output_dir / "research_paper_siddarth_boggarapu.txt"
    with open(txt_path, "w") as f:
        f.write("Fallback report due to missing python-docx or other error.\n\n")
        f.write(summary.to_string())
    print("DOCX generation failed:", e)

# ZIP everything
zip_path = output_dir / "propeller_analysis_package.zip"
with zipfile.ZipFile(zip_path, "w") as zf:
    for p in plots:
        zf.write(p, arcname=p.name)
    zf.write(clean_csv, arcname=clean_csv.name)
    zf.write(summary_csv, arcname=summary_csv.name)
    if docx_created:
        zf.write(docx_path, arcname=docx_path.name)
    else:
        zf.write(txt_path, arcname=txt_path.name)

print("Done. Outputs written to:", output_dir)
print("ZIP:", zip_path)
